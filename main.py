import os 
import datetime as dt
import subprocess
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox 
import docx 


class InvoiceGen:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title('InvoiceGen 1.1')
        self.root.geometry('500x600')


        self.partner_label = tk.Label(self.root, text='Partner')
        self.partner_street_label = tk.Label(self.root, text='Partner Street')
        self.partner_zip_city_country_label = tk.Label(self.root, text='Partner Zip_City_Country')
        self.invoice_number_label = tk.Label(self.root, text='Invoice Number')
        self.service_description_label = tk.Label(self.root, text='Service Description')
        self.service_amount_label = tk.Label(self.root, text='Amount')
        self.service_single_price_label = tk.Label(self.root, text='Single Price')
        self.payment_method_label = tk.Label(self.root, text='Payment Method')


        self.payment_options = {

            'Main Bank':{
                'Recipient': 'Jeff Muney',
                'Bank': 'PesaMingi',
                'Branch':'KonaDongo',
                'Account':'1223 4578 8900 7367'    
            },

            'MPESA':{
                'Recipient': 'Pita Nuako',
                'Bank': 'MPESA',
                'Branch':'NA',
                'Account':'NAMBA YANGU'    
            },

            'Second Bank':{
                'Recipient': 'Bestprof',
                'Bank': 'Manibank',
                'Branch':'Nairobe',
                'Account':'1223 4578 8900 7367'    
            }
        }


        self.partner_entry = tk.Entry(self.root)
        self.partner_street_entry = tk.Entry(self.root)
        self.partner_zip_city_country_entry = tk.Entry(self.root)
        self.invoice_number_entry = tk.Entry(self.root)
        self.service_description_entry = tk.Entry(self.root)
        self.service_amount_entry = tk.Entry(self.root)
        self.service_single_price_entry = tk.Entry(self.root)
        self.payment_method_entry = tk.Entry(self.root)


        self.payment_option = tk.StringVar(self.root)
        self.payment_option.set('Main Bank')

        self.payment_method_dropdown = tk.OptionMenu(self.root, self.payment_option, "Main Bank","MPESA","Second Bank")

        self.make_button = tk.Button(self.root, text='Generate Invoice', command=self.generate_invoice)

        padding_options={'fill':'x', 'expand':True, 'padx':5, 'pady': 2}


        self.partner_label.pack(padding_options)
        self.partner_entry.pack(padding_options)
        
        self.partner_street_label.pack(padding_options)
        self.partner_street_entry.pack(padding_options)

        self.partner_zip_city_country_label.pack(padding_options)
        self.partner_zip_city_country_entry.pack(padding_options)

        self.invoice_number_label.pack(padding_options)
        self.invoice_number_entry.pack(padding_options)

        self.service_description_label.pack(padding_options)
        self.service_description_entry.pack(padding_options)

        self.service_amount_label.pack(padding_options)
        self.service_amount_entry.pack(padding_options)

        self.service_single_price_label.pack(padding_options)
        self.service_single_price_entry.pack(padding_options)
        
        self.payment_method_label.pack(padding_options)
        self.payment_method_dropdown.pack(padding_options)
        self.make_button.pack(padding_options)
        #self.payment_method_entry = tk.Entry(self.root)

        self.root.mainloop()
    @staticmethod
    def write_doc(paragraph, old_text, new_text):
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)


    def generate_invoice(self):
        template = docx.Document(r'Template.docx')
        selected_payment_method = self.payment_options[self.payment_option.get()]
        
        try:

            replacements ={

                "[Date]": dt.datetime.today().strftime('%Y-%m-%d'),
                "[Partner]": self.partner_entry.get(),
                "[Partner Street]": self.partner_street_entry.get(),
                "[Partner Zip_City_Country]": self.partner_zip_city_country_entry.get(),
                "[Invoice Number]": self.invoice_number_entry.get(),
                "[Service Description]": self.service_description_entry.get(),
                "[Amount]": self.service_amount_entry.get(),
                "[Single Price]": f"${float(self.service_single_price_entry.get()):.2f}",
                "[Full Price]": f'${float(self.service_amount_entry.get())*float(self.service_single_price_entry.get()):2f}',
                "[Recipient]": selected_payment_method['Recipient'],
                "[Bank]": selected_payment_method['Bank'],
                "[Branch]": selected_payment_method['Branch'],
                "[Account]": selected_payment_method['Branch'],
                



            }
        except ValueError: 
            messagebox.showerror('Error', 'Invalid amount or price!')
            return

        for paragraph in list(template.paragraphs):
            for old_text, new_text in replacements.items():
                self.write_doc(paragraph, old_text, new_text)

        for table in template.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            self.write_doc(paragraph, old_text, new_text)

        save_path = filedialog.asksaveasfilename(defaultextension='pdf', filetypes=[('PDF documents', '*.pdf')])
        template.save('Template.docx')

        import docx2pdf
        from docx2pdf import convert
        convert("Template.docx", save_path)
        messagebox.showinfo('Success', 'Invoice created and saved successfully!')

if __name__ == '__main__':
    InvoiceGen()

